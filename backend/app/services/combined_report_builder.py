"""
Combined Project Report -- takes any set of past calculator runs (batch
matrix, pile capacity, pile group, rock bearing, retaining wall, liquefaction,
lateral capacity, rock-socketed pile, ground improvement -- whatever the
person picks) and assembles them into ONE downloadable DOCX. Added 14 Aug
2026, per Raahi's request to have every calculator "connected" so a final
report can be produced from whichever combination of results are relevant to
a given project.

This is deliberately DIFFERENT from report_builder.py (the existing "Auto
Report Generation" feature, which is scoped to exactly one borehole chart +
one batch result + an AI summary). This module works from CalculationLog
rows -- every calculator already writes one there on every run (see each
endpoint in routers/calculators.py) -- so "connecting" the calculators meant
building a way to READ those logs back and combine any subset of them, not
changing how the calculators themselves work.

Two levels of detail per calculator type, by design:
  1. Calculators this module "knows" well (built/verified in this same
     project): pile_capacity, pile_group_analysis, batch_matrix get a
     hand-built section (headline numbers highlighted, compact result
     table matching what their own pages show).
  2. Everything else (rock_bearing_capacity, retaining_wall_stability,
     liquefaction_analysis, lateral_capacity, rock_socket_pile,
     ground_improvement, and any future calculator_type) gets a GENERIC
     section: every top-level scalar field in its result, auto-tabulated.
     This is deliberate -- guessing at exact field names for calculators
     not touched in this session risks silently showing wrong/stale labels,
     which is worse than an honest auto-generated table. Full per-type
     hand-built sections for the rest can be added later, one at a time, as
     each one gets used for a real report and any gaps show up.

NOT covered (flagged, not silently dropped):
  - PDF export of this combined report (DOCX only, same as the other
    auto-report feature).
  - Full layer-by-layer working tables (skin friction segments, settlement
    sub-layers, batch combinations beyond the results table) are NOT
    reproduced here -- only headline/summary figures. The in-app calculator
    pages remain the place to check full working; this report is the
    hand-off document.
  - Editing inside the app -- one-shot DOCX download, edit in Word after.
"""
import io
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

CALC_TYPE_TITLES = {
    "pile_capacity": "Single Pile Capacity",
    "pile_group_analysis": "Pile Group Analysis",
    "rock_bearing_capacity": "Bearing Capacity on Rock",
    "retaining_wall_stability": "Retaining Wall Stability",
    "batch_matrix": "Batch Foundation Analysis (Shear + Settlement SBC)",
    "liquefaction_analysis": "Liquefaction Analysis",
    "lateral_capacity": "Lateral Pile Capacity",
    "rock_socket_pile": "Rock-Socketed Pile Capacity",
    "ground_improvement": "Ground Improvement",
}

# Keys that are either huge (full layer-by-layer working, meant for the
# in-app page, not this summary report), redundant with the section
# heading/metadata already shown, or not human-readable on their own.
_SKIP_RESULT_KEYS = {
    "layer_report", "positions", "pile_positions_m", "end_bearing_candidates",
    "sub_layers", "combinations", "layers_used", "estimated_fields",
    "borehole_id", "cap_load_distribution", "settlement", "block_failure",
    "single_pile", "group_efficiency", "group_capacity_efficiency_method",
    "group_capacity_block_method",
}
_SKIP_INPUT_KEYS = {"borehole_id", "overrides", "widths_m", "depths_m"}


def _headline(calc_type: str, result: dict) -> str:
    """One-line summary for a calculation -- used both in the history list
    (so the person can tell runs apart without opening full JSON) and fed to
    the AI as context for the overall conclusion. Only uses fields this
    module KNOWS exist (hand-built types); falls back to a generic 'ran
    successfully' line for everything else rather than guessing field names."""
    try:
        if calc_type == "pile_capacity":
            return f"Allowable compression {result.get('allowable_compression_capacity_t')} t, uplift {result.get('allowable_uplift_capacity_t')} t"
        if calc_type == "pile_group_analysis":
            return f"{result.get('layout')} piles -- governing allowable {result.get('governing_group_capacity_t')} t ({result.get('governing_mode')})"
        if calc_type == "batch_matrix":
            crit = result.get("critical_combination") or {}
            return f"{result.get('successful')}/{result.get('total')} combinations -- critical {crit.get('width_m')}m x {crit.get('depth_m')}m = {crit.get('recommended_sbc')} t/m2"
    except Exception:
        pass
    return "See results table below."


def _scalar_rows(d: dict, exclude: set) -> list[tuple[str, str]]:
    """Every top-level scalar (number/string/bool) field in d, plus one
    level of nested scalars (dotted key), excluding the given key names.
    This is the generic fallback -- honest about showing whatever is
    actually in the result rather than a hand-picked (and possibly stale)
    subset."""
    rows = []
    for k, v in d.items():
        if k in exclude or v is None or v == "":
            continue
        if isinstance(v, bool) or isinstance(v, (int, float, str)):
            rows.append((k.replace("_", " "), str(v)))
        elif isinstance(v, dict):
            for k2, v2 in v.items():
                if v2 is None or v2 == "":
                    continue
                if isinstance(v2, bool) or isinstance(v2, (int, float, str)):
                    rows.append((f"{k}.{k2}".replace("_", " "), str(v2)))
    return rows


def _add_kv_table(doc, rows: list[tuple[str, object]]):
    rows = [(label, value) for label, value in rows if value is not None and value != ""]
    if not rows:
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)


def _add_warnings(doc, result: dict):
    warnings = result.get("warnings")
    if warnings:
        p = doc.add_paragraph()
        p.add_run("Assumptions & warnings:").italic = True
        for w in warnings[:8]:  # cap -- this is a summary report, not the full in-app list
            doc.add_paragraph(str(w), style="List Bullet")


def _add_batch_matrix_section(doc, inputs: dict, result: dict):
    doc.add_paragraph(f"Successful combinations: {result.get('successful')} / {result.get('total')}")
    combos = [c for c in result.get("combinations", []) if "error" not in c]
    if combos:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Width (m)", "Depth (m)", "Shear SBC", "Settlement SBC", "Recommended SBC", "Governing"]):
            hdr[i].text = h
        for c in combos:
            row = table.add_row().cells
            row[0].text = str(c.get("width_m", ""))
            row[1].text = str(c.get("depth_m", ""))
            row[2].text = str(c.get("shear_sbc", ""))
            row[3].text = str(c.get("settlement_sbc", ""))
            row[4].text = str(c.get("recommended_sbc", ""))
            row[5].text = str(c.get("governing", ""))
    critical = result.get("critical_combination")
    if critical:
        p = doc.add_paragraph()
        p.add_run("Critical combination (governing, lowest recommended SBC): ").bold = True
        p.add_run(
            f"Width {critical.get('width_m')}m, Depth {critical.get('depth_m')}m -- "
            f"Recommended SBC = {critical.get('recommended_sbc')} t/m2 ({critical.get('governing')})"
        )


def _add_pile_capacity_section(doc, inputs: dict, result: dict):
    dia_mm = round((inputs.get("diameter_m") or 0) * 1000)
    p = doc.add_paragraph()
    p.add_run(f"Pile: \u00d8{dia_mm}mm x {inputs.get('pile_length_m')}m below cutoff, code {result.get('code', inputs.get('code', ''))}. ").bold = True
    _add_kv_table(doc, [
        ("Ultimate compression capacity (t)", result.get("ultimate_compression_capacity_t")),
        ("Allowable compression capacity (t)", result.get("allowable_compression_capacity_t")),
        ("Ultimate uplift capacity (t)", result.get("ultimate_uplift_capacity_t")),
        ("Allowable uplift capacity (t)", result.get("allowable_uplift_capacity_t")),
    ])
    _add_warnings(doc, result)


def _add_pile_group_section(doc, inputs: dict, result: dict):
    p = doc.add_paragraph()
    p.add_run(f"Group: {result.get('layout')} piles, envelope {result.get('group_length_m')}m x {result.get('group_width_m')}m. ").bold = True
    _add_kv_table(doc, [
        ("Single pile allowable (t)", result.get("single_pile", {}).get("allowable_compression_capacity_t")),
        ("Group allowable -- efficiency method (t)", result.get("group_capacity_efficiency_method", {}).get("allowable_t")),
        ("Group allowable -- block failure method (t)", result.get("group_capacity_block_method", {}).get("allowable_t")),
        ("Governing group capacity (t)", result.get("governing_group_capacity_t")),
        ("Governing mode", result.get("governing_mode")),
        ("Max pile load (t)", result.get("cap_load_distribution", {}).get("max_pile_load_t")),
        ("Allowable per pile, efficiency-reduced (t)", result.get("cap_load_distribution", {}).get("allowable_per_pile_t")),
    ])
    settlement = result.get("settlement")
    if settlement:
        _add_kv_table(doc, [
            ("Settlement (mm)", settlement.get("result")),
            ("Raft depth (m)", settlement.get("raft_depth_m")),
        ])
    _add_warnings(doc, result)


_SECTION_BUILDERS = {
    "batch_matrix": _add_batch_matrix_section,
    "pile_capacity": _add_pile_capacity_section,
    "pile_group_analysis": _add_pile_group_section,
}


def _add_generic_section(doc, calc_type: str, inputs: dict, result: dict):
    input_rows = _scalar_rows(inputs, _SKIP_INPUT_KEYS)
    if input_rows:
        doc.add_paragraph("Inputs:").runs[0].italic = True
        _add_kv_table(doc, input_rows)
    result_rows = _scalar_rows(result, _SKIP_RESULT_KEYS)
    if result_rows:
        doc.add_paragraph("Results:").runs[0].italic = True
        _add_kv_table(doc, result_rows)
    if not input_rows and not result_rows:
        doc.add_paragraph("(No summarizable scalar fields found in this result -- see the in-app calculator for full detail.)")
    _add_warnings(doc, result)


def build_combined_report_docx(
    entries: list[dict], title: str, project_name: str | None, site_location: str | None,
    ai_summary: str | None,
) -> bytes:
    """entries: list of {calculator_type, created_at, inputs, result}, in the
    order the person picked them. Builds one section per entry, an "Included
    Calculations" index at the top, and an optional AI-written overall
    conclusion at the end."""
    doc = DocxDocument()

    t = doc.add_heading(title or "Combined Geotechnical Engineering Report", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    if project_name:
        meta.add_run(f"Project: {project_name}").bold = True
    if site_location:
        meta.add_run(f"   |   Location: {site_location}")
    if meta.runs:
        meta.runs[0].font.size = Pt(11)

    doc.add_heading("Included Calculations", level=1)
    for e in entries:
        title_label = CALC_TYPE_TITLES.get(e["calculator_type"], e["calculator_type"])
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{title_label} ").bold = True
        p.add_run(f"-- {_headline(e['calculator_type'], e['result'])}")

    for e in entries:
        calc_type = e["calculator_type"]
        title_label = CALC_TYPE_TITLES.get(calc_type, calc_type)
        doc.add_heading(title_label, level=1)
        if e.get("created_at"):
            ts = doc.add_paragraph(f"Run on: {e['created_at']}")
            ts.runs[0].font.size = Pt(8)
            ts.runs[0].font.italic = True
        builder = _SECTION_BUILDERS.get(calc_type)
        if builder:
            builder(doc, e["inputs"], e["result"])
        else:
            _add_generic_section(doc, calc_type, e["inputs"], e["result"])

    if ai_summary:
        doc.add_heading("Overall Engineering Conclusion", level=1)
        for para in ai_summary.split("\n"):
            if para.strip():
                doc.add_paragraph(para)

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "Generated by RaahiGeo, combining results already run in-app -- verify all figures "
        "and check the full per-calculator working in-app before use in a submitted design."
    )
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
