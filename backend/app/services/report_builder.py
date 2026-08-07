"""
Auto-report generation -- combines a borehole log chart, batch analysis
results table, and an AI-written summary into one downloadable DOCX.

Added 7 Aug 2026 -- this is the roadmap item Raahi flagged as "the next
real open item" back when this project's status doc was first written.
Deliberately scoped to the exact three pieces Raahi asked for (borehole
log chart + batch results + summary combined into one report) rather than
trying to build a full customizable report builder in one pass -- the
existing manual section-by-section Reports page (routers/reports.py's
/generate + /export/docx) still exists separately for anything more
freeform.

NOT covered here (same "flag it, don't fake it" policy as every other
calculator in this app):
  - PDF export of this combined report (DOCX only for now -- the existing
    /export/pdf endpoint is for the separate manual-sections flow and
    doesn't know about charts/tables).
  - Multiple boreholes / multiple batch runs in one report -- one borehole,
    one batch result per report.
  - Any editing of the generated report inside the app -- it's a one-shot
    DOCX download, open and edit in Word after.
"""
import io
import matplotlib
matplotlib.use("Agg")  # headless -- no display server on Render
import matplotlib.pyplot as plt
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_borehole_chart_png(borehole) -> bytes:
    """
    Simple depth-vs-N-value strip log: SPT N-value as a line/marker plot,
    each layer's USCS classification annotated alongside, water table shown
    as a dashed line. Depth increases downward (inverted y-axis), matching
    how a borehole log is conventionally drawn.
    """
    layers = sorted(borehole.layers, key=lambda l: l.from_m)
    max_depth = max((l.to_m for l in layers), default=1.0)

    fig, (ax_log, ax_n) = plt.subplots(
        1, 2, figsize=(6, max(4, max_depth * 0.35)), sharey=True,
        gridspec_kw={"width_ratios": [1.3, 1]},
    )

    # -- Left panel: classification strip --
    for l in layers:
        ax_log.barh(
            y=(l.from_m + l.to_m) / 2, width=1, height=(l.to_m - l.from_m),
            left=0, color="#94a3b8", alpha=0.25, edgecolor="#475569", linewidth=0.6,
        )
        label = l.classification or (l.description or "")[:18] or "—"
        ax_log.text(0.5, (l.from_m + l.to_m) / 2, label, ha="center", va="center", fontsize=7.5, color="#1e293b")
    ax_log.set_xlim(0, 1)
    ax_log.set_xticks([])
    ax_log.set_ylabel("Depth (m)")
    ax_log.set_title("Strata", fontsize=9)
    ax_log.invert_yaxis()

    # -- Right panel: N-value --
    n_depths = [(l.from_m + l.to_m) / 2 for l in layers if l.n_value is not None]
    n_values = [l.n_value for l in layers if l.n_value is not None]
    if n_values:
        ax_n.plot(n_values, n_depths, marker="o", markersize=4, color="#7c3aed", linewidth=1.3)
    ax_n.set_xlabel("N-value")
    ax_n.set_title("SPT N", fontsize=9)
    ax_n.grid(True, alpha=0.3)

    if borehole.water_table_depth_m is not None:
        for ax in (ax_log, ax_n):
            ax.axhline(borehole.water_table_depth_m, color="#0ea5e9", linestyle="--", linewidth=1)
        ax_n.text(
            ax_n.get_xlim()[1] * 0.5, borehole.water_table_depth_m, " WT",
            color="#0ea5e9", fontsize=7.5, va="bottom",
        )

    fig.suptitle(f"Borehole {borehole.borehole_id}", fontsize=11, fontweight="bold")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def build_batch_report_docx(borehole, batch_result: dict, ai_summary: str) -> bytes:
    """
    Combines the three pieces into one DOCX: title/project info, the
    borehole log chart, the batch results table (governing SBC per
    width/depth combination), the critical (lowest recommended SBC)
    combination called out, and the AI-written summary paragraph.
    """
    doc = DocxDocument()

    title = doc.add_heading(f"Batch Foundation Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"Borehole: {borehole.borehole_id}").bold = True
    if borehole.project_name:
        meta.add_run(f"   |   Project: {borehole.project_name}")
    meta_style = meta.runs[0].font
    meta_style.size = Pt(11)

    # -- Borehole log chart --
    doc.add_heading("Borehole Log", level=1)
    chart_png = generate_borehole_chart_png(borehole)
    doc.add_picture(io.BytesIO(chart_png), width=Inches(5.5))

    # -- Batch results table --
    doc.add_heading("Batch Analysis Results", level=1)
    combos = [c for c in batch_result.get("combinations", []) if "error" not in c]
    if combos:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Width (m)", "Depth (m)", "Shear SBC", "Settlement SBC", "Recommended SBC", "Governing"]):
            hdr[i].text = h
        for c in combos:
            row = table.add_row().cells
            row[0].text = str(c.get("width_m", ""))
            row[1].text = str(c.get("depth_m", ""))
            row[2].text = str(c.get("shear_sbc", ""))
            row[3].text = str(c.get("settlement_sbc", ""))
            row[4].text = str(c.get("recommended_sbc", ""))
            row[5].text = str(c.get("governing", ""))
    else:
        doc.add_paragraph("No successful combinations in this batch run.")

    critical = batch_result.get("critical_combination")
    if critical:
        p = doc.add_paragraph()
        p.add_run("Critical combination (governing, lowest recommended SBC): ").bold = True
        p.add_run(
            f"Width {critical.get('width_m')}m, Depth {critical.get('depth_m')}m -- "
            f"Recommended SBC = {critical.get('recommended_sbc')} t/m² ({critical.get('governing')})"
        )

    # -- AI summary --
    doc.add_heading("Summary", level=1)
    for para in ai_summary.split("\n"):
        if para.strip():
            doc.add_paragraph(para)

    doc.add_paragraph()
    footer = doc.add_paragraph("Generated by RaahiGeo -- verify all figures before use in a submitted design.")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
