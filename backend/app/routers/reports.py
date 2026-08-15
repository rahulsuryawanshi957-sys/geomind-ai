import io
import json
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from app.schemas import ReportSectionRequest, AutoReportRequest, CombinedReportRequest
from app.rag.retrieval import retrieve
from app.services.llm import generate_report_section
from app.services.report_builder import build_batch_report_docx
from app.services.combined_report_builder import build_combined_report_docx, _headline
from app.database import get_db
from app.models import BoreholeProfile, CalculationLog

router = APIRouter(prefix="/api/reports", tags=["reports"])

SECTION_TYPES = [
    "SBC Recommendation", "Foundation Recommendation", "Pile Recommendation",
    "Liquefaction Summary", "Settlement Summary", "Design Notes", "Engineering Conclusion",
]


@router.get("/section-types")
def section_types():
    return SECTION_TYPES


@router.post("/generate")
def generate_section(req: ReportSectionRequest):
    query = req.reference_query or req.section_type
    chunks = retrieve(query)
    content = generate_report_section(req.section_type, req.project_inputs, chunks)
    return {"section_type": req.section_type, "content": content, "sources_used": len(chunks)}


@router.post("/export/docx")
def export_docx(sections: dict):
    """sections: {"Project Title": "...", "sections": [{"title": "...", "content": "..."}]}"""
    doc = DocxDocument()
    doc.add_heading(sections.get("title", "Geotechnical Report"), level=0)
    for sec in sections.get("sections", []):
        doc.add_heading(sec["title"], level=1)
        for para in sec["content"].split("\n"):
            if para.strip():
                doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=raahigeo_report.docx"},
    )


@router.post("/export/pdf")
def export_pdf(sections: dict):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph(sections.get("title", "Geotechnical Report"), styles["Title"]), Spacer(1, 12)]
    for sec in sections.get("sections", []):
        story.append(Paragraph(sec["title"], styles["Heading2"]))
        for para in sec["content"].split("\n"):
            if para.strip():
                story.append(Paragraph(para, styles["BodyText"]))
        story.append(Spacer(1, 10))
    doc.build(story)
    buf.seek(0)
    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=raahigeo_report.pdf"},
    )


@router.post("/auto-generate")
def auto_generate(req: AutoReportRequest, db: Session = Depends(get_db)):
    """
    Combines a borehole log chart, the given batch analysis results table,
    and an AI-written summary into one downloadable DOCX -- see
    app/services/report_builder.py for exactly what's covered (DOCX only,
    one borehole + one batch result per report; no PDF export of this
    combined report yet).
    """
    borehole = db.query(BoreholeProfile).filter(BoreholeProfile.id == req.borehole_id).first()
    if not borehole:
        raise HTTPException(404, f"Borehole '{req.borehole_id}' not found.")
    if not borehole.layers:
        raise HTTPException(422, f"Borehole '{req.borehole_id}' has no soil layers -- nothing to chart.")

    critical = req.batch_result.get("critical_combination") or {}
    summary_inputs = {
        "borehole_id": borehole.borehole_id,
        "project_name": borehole.project_name,
        "total_combinations_run": req.batch_result.get("total"),
        "successful_combinations": req.batch_result.get("successful"),
        "critical_width_m": critical.get("width_m"),
        "critical_depth_m": critical.get("depth_m"),
        "critical_recommended_sbc_t_m2": critical.get("recommended_sbc"),
        "critical_governing": critical.get("governing"),
    }
    chunks = retrieve(f"foundation recommendation SBC {borehole.project_name or borehole.borehole_id}")
    ai_summary = generate_report_section("Batch Analysis Summary", summary_inputs, chunks)

    docx_bytes = build_batch_report_docx(borehole, req.batch_result, ai_summary)
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=raahigeo_batch_report_{borehole.borehole_id}.docx"},
    )


@router.post("/combined-generate")
def combined_generate(req: CombinedReportRequest, db: Session = Depends(get_db)):
    """Combined Project Report -- picks any set of past calculator runs
    (by CalculationLog id) and assembles them into one DOCX. See
    combined_report_builder.py's module docstring for exactly how each
    calculator type is rendered (hand-built for pile_capacity/
    pile_group_analysis/batch_matrix, generic auto-table for everything
    else)."""
    if not req.log_ids:
        raise HTTPException(422, "Select at least one calculation to include.")
    logs = db.query(CalculationLog).filter(CalculationLog.id.in_(req.log_ids)).all()
    if not logs:
        raise HTTPException(404, "None of the selected calculations were found -- they may have been logged before a database reset.")
    logs_by_id = {l.id: l for l in logs}

    entries = []
    for lid in req.log_ids:  # preserve the order the person picked them in
        log = logs_by_id.get(lid)
        if not log:
            continue
        try:
            inputs = json.loads(log.inputs_json) if log.inputs_json else {}
        except Exception:
            inputs = {}
        try:
            result = json.loads(log.result_json) if log.result_json else {}
        except Exception:
            result = {}
        entries.append({"calculator_type": log.calculator_type, "created_at": str(log.created_at), "inputs": inputs, "result": result})

    if not entries:
        raise HTTPException(404, "None of the selected calculations could be loaded.")

    ai_summary = None
    if req.write_ai_summary:
        digest = {
            "project_name": req.project_name,
            "site_location": req.site_location,
            "calculations_included": [
                {"type": e["calculator_type"], "headline": _headline(e["calculator_type"], e["result"])}
                for e in entries
            ],
        }
        chunks = retrieve(f"engineering conclusion foundation recommendation {req.project_name or ''}")
        ai_summary = generate_report_section("Engineering Conclusion", digest, chunks)

    docx_bytes = build_combined_report_docx(entries, req.title, req.project_name, req.site_location, ai_summary)
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=raahigeo_combined_report.docx"},
    )
